"""Convert CHAPTERS_FOUR_AND_FIVE.md to a formatted academic .docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
DOUBLE_SPACING = 2.0


def set_run_font(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def configure_paragraph(paragraph, *, align=None, bold: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = DOUBLE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    if bold and paragraph.runs:
        for run in paragraph.runs:
            run.bold = True


def add_formatted_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    align=None,
    heading_bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    configure_paragraph(paragraph, align=align)

    if not text:
        return

    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, bold=heading_bold)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            configure_paragraph(paragraph)
            run = paragraph.add_run(cell_text)
            set_run_font(run, bold=(r_idx == 0))
            for extra_p in cell.paragraphs[1:]:
                extra_p._element.getparent().remove(extra_p._element)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = DOUBLE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_defaults(doc)

    i = 0
    pending_list: list[tuple[str, str]] = []

    def flush_list() -> None:
        nonlocal pending_list
        for kind, item in pending_list:
            style = "List Number" if kind == "num" else "List Bullet"
            paragraph = doc.add_paragraph(style=style)
            configure_paragraph(paragraph)
            add_inline_runs(paragraph, item)
        pending_list = []

    def add_inline_runs(paragraph, text: str) -> None:
        pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                set_run_font(run, bold=True)
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                set_run_font(run, italic=True)
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                set_run_font(run)
            else:
                run = paragraph.add_run(part)
                set_run_font(run)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_list()
            i += 1
            continue

        if stripped == "---":
            flush_list()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_list()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    table_lines.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            flush_list()
            title = stripped[2:].strip()
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
            run = paragraph.add_run(title.upper())
            set_run_font(run, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_list()
            heading = stripped[3:].strip()
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph)
            run = paragraph.add_run(heading)
            set_run_font(run, bold=True)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_list()
            heading = stripped[4:].strip()
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph)
            run = paragraph.add_run(heading)
            set_run_font(run, bold=True)
            i += 1
            continue

        if stripped.startswith("#### "):
            flush_list()
            heading = stripped[5:].strip()
            paragraph = doc.add_paragraph()
            configure_paragraph(paragraph)
            run = paragraph.add_run(heading)
            set_run_font(run, bold=True)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            pending_list.append(("bullet", bullet_match.group(1)))
            i += 1
            continue

        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            pending_list.append(("num", num_match.group(1)))
            i += 1
            continue

        flush_list()
        paragraph = doc.add_paragraph()
        configure_paragraph(paragraph)
        add_inline_runs(paragraph, stripped)
        i += 1

    flush_list()
    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md_file = root / "CHAPTERS_FOUR_AND_FIVE.md"
    out_file = root / "CHAPTERS_FOUR_AND_FIVE.docx"
    convert_markdown_to_docx(md_file, out_file)
    print(f"Created: {out_file}")